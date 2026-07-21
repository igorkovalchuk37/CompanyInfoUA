from __future__ import annotations

import base64
import copy
import io
import json
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from openai import OpenAI

from ua_business_map_release.render_business_map import render as render_business_map


REPORT_DATE = date.today().isoformat()


@dataclass(frozen=True)
class ReportConfig:
    openai_model: str = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")
    openai_web_search_tool_type: str = os.getenv("OPENAI_WEB_SEARCH_TOOL_TYPE", "web_search")
    openai_reasoning_effort: str = os.getenv("OPENAI_REASONING_EFFORT", "medium")
    allow_live_openai: bool = os.getenv("ALLOW_LIVE_OPENAI", "").lower() in {"1", "true", "yes", "on"}


def _normalize_text(value: Any, fallback: str = "не встановлено") -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text or fallback


def _normalize_list(items: Any) -> list[dict[str, Any]]:
    if not isinstance(items, list):
        return []
    return [item for item in items if isinstance(item, dict)]


def _extract_json_object(text: str) -> dict[str, Any]:
    text = text.strip()
    if not text:
        raise ValueError("OpenAI response is empty")

    match = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not match:
        raise ValueError("OpenAI response does not contain JSON")
    return json.loads(match.group(0))


def _replace_paragraph_text(paragraph: Any, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _iter_all_paragraphs(document: Document) -> list[Any]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _replace_tokens(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in _iter_all_paragraphs(document):
        combined = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        updated = combined
        for token, value in replacements.items():
            updated = updated.replace(token, value)
        if updated != combined:
            _replace_paragraph_text(paragraph, updated)


def _find_row_index(table: Any, token: str) -> int:
    for idx, row in enumerate(table.rows):
        row_text = " | ".join(cell.text for cell in row.cells)
        if token in row_text:
            return idx
    raise ValueError(f"Template row with token {token} not found")


def _fill_repeating_table(table: Any, token: str, rows_data: list[dict[str, str]]) -> None:
    row_index = _find_row_index(table, token)
    template_row = table.rows[row_index]
    template_tr = template_row._tr

    for row_data in rows_data:
        new_tr = copy.deepcopy(template_tr)
        table._tbl.insert(row_index, new_tr)
        new_row = table.rows[row_index]
        row_replacements = {key: _normalize_text(value, "") for key, value in row_data.items()}
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                combined = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
                updated = combined
                for placeholder, value in row_replacements.items():
                    updated = updated.replace(placeholder, value)
                if updated != combined:
                    _replace_paragraph_text(paragraph, updated)
        row_index += 1

    table._tbl.remove(template_tr)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)

    run.append(run_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _insert_map_image(document: Document, image_bytes: bytes) -> None:
    for paragraph in document.paragraphs:
        if "{{MAP_IMAGE}}" in paragraph.text:
            _replace_paragraph_text(paragraph, "")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(image_bytes), width=Cm(16.0))
            return
    raise ValueError("Map placeholder {{MAP_IMAGE}} not found in template")


def _find_table_after_heading(document: Document, heading_text: str) -> Any:
    body = document._body._body
    found_heading = False
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t")).strip()
            if text == heading_text:
                found_heading = True
        elif child.tag.endswith("}tbl") and found_heading:
            for table in document.tables:
                if table._tbl == child:
                    return table
    raise ValueError(f"Table after heading '{heading_text}' not found")


def _fallback_beneficiaries() -> list[dict[str, str]]:
    return [{
        "{{BENEFICIARY_PERSON}}": "не встановлено",
        "{{BENEFICIARY_STATUS}}": "",
        "{{CONTROL_BASIS}}": "",
        "{{RELIABILITY}}": "",
    }]


def _fallback_events() -> list[dict[str, str]]:
    return [{
        "{{EVENT_OR_RISK}}": "не встановлено",
        "{{IMPACT}}": "",
        "{{DATE_OR_PERIOD}}": "",
    }]


def _fallback_sources() -> list[dict[str, str]]:
    return [{
        "{{SOURCE_NAME}}": "не встановлено",
        "{{SOURCE_DATE}}": "",
        "{{SOURCE_FACT}}": "",
        "{{SOURCE_URL}}": "",
    }]


def _normalize_report_payload(payload: dict[str, Any], edrpou: str) -> dict[str, Any]:
    company = payload.get("company") if isinstance(payload.get("company"), dict) else {}
    group = payload.get("group") if isinstance(payload.get("group"), dict) else {}

    beneficiaries = _normalize_list(payload.get("beneficiaries"))
    events = _normalize_list(payload.get("events"))
    sources = _normalize_list(payload.get("sources"))
    map_data = payload.get("map") if isinstance(payload.get("map"), dict) else {}

    return {
        "report_date": _normalize_text(payload.get("report_date"), REPORT_DATE),
        "company": {
            "company_name": _normalize_text(company.get("company_name")),
            "full_name": _normalize_text(company.get("full_name")),
            "status": _normalize_text(company.get("status")),
            "registration_date": _normalize_text(company.get("registration_date")),
            "legal_address": _normalize_text(company.get("legal_address")),
            "director": _normalize_text(company.get("director")),
            "main_kved": _normalize_text(company.get("main_kved")),
            "activity_and_assets": _normalize_text(company.get("activity_and_assets")),
            "edrpou": edrpou,
        },
        "group": {
            "membership_and_role": _normalize_text(group.get("membership_and_role")),
            "group_name": _normalize_text(group.get("group_name")),
            "controlling_persons": _normalize_text(group.get("controlling_persons")),
            "business_lines": _normalize_text(group.get("business_lines")),
            "geography": _normalize_text(group.get("geography")),
            "asset_types": _normalize_text(group.get("asset_types")),
            "key_group_companies": _normalize_text(group.get("key_group_companies")),
        },
        "beneficiaries": beneficiaries,
        "events": events,
        "sources": sources,
        "map": map_data,
    }


def build_mock_report_data(edrpou: str, company_name_hint: str | None = None) -> dict[str, Any]:
    company_name = company_name_hint or f"Тестова компанія {edrpou}"
    return {
        "report_date": REPORT_DATE,
        "company": {
            "company_name": company_name,
            "full_name": f"ТОВ \"{company_name}\"",
            "status": "зареєстровано",
            "registration_date": "не встановлено",
            "legal_address": "не встановлено",
            "director": "не встановлено",
            "main_kved": "не встановлено",
            "activity_and_assets": (
                "Тестовий безкоштовний режим. Реальні дані OpenAI не викликались. "
                "Цей документ призначений для перевірки Render, Swagger, карти та збірки DOCX."
            ),
            "edrpou": edrpou,
        },
        "group": {
            "membership_and_role": (
                "Тестовий режим. Належність до бізнес-групи не досліджувалась, "
                "оскільки для цього запуску OpenAI API не використовувався."
            ),
            "group_name": "Тестова бізнес-група",
            "controlling_persons": "не встановлено",
            "business_lines": "Тестова генерація документа",
            "geography": "Україна",
            "asset_types": "hq, production",
            "key_group_companies": f"{company_name}; Тест Логістик; Тест Виробництво",
        },
        "beneficiaries": [
            {
                "person": "не встановлено",
                "status": "",
                "control_basis": "",
                "reliability": "",
            }
        ],
        "events": [
            {
                "event_or_risk": "Тестовий запуск сервісу",
                "impact": "Перевірка генерації документа без витрат на OpenAI API",
                "date_or_period": REPORT_DATE,
            }
        ],
        "sources": [
            {
                "source_name": "Локальний mock-режим сервісу",
                "source_date": REPORT_DATE,
                "source_fact": "Документ згенеровано без реального дослідження зовнішніх джерел",
                "source_url": "https://companyinfoua.onrender.com/docs",
            }
        ],
        "map": {
            "group_name": "Тестова бізнес-група",
            "description": "Тестова карта активів",
            "show_title": False,
            "show_legend": False,
            "derive_active_regions": True,
            "assets": [
                {
                    "name": "Тестовий офіс",
                    "type": "hq",
                    "region_id": "UA32",
                    "city": "Київ",
                    "show_label": True,
                },
                {
                    "name": "Тестовий майданчик",
                    "type": "production",
                    "region_id": "UA12",
                    "city": "Дніпро",
                    "show_label": True,
                },
            ],
        },
    }


def _research_prompt(edrpou: str, company_name_hint: str | None, additional_instructions: str | None) -> str:
    hint_line = f"Назва-підказка: {company_name_hint}\n" if company_name_hint else ""
    extra_line = f"Додаткові інструкції користувача: {additional_instructions}\n" if additional_instructions else ""
    return f"""
Підготуй структуровані дані для довідки по українській компанії за кодом ЄДРПОУ {edrpou}.
{hint_line}{extra_line}
Працюй українською мовою.
Дата формування звіту: {REPORT_DATE}.

Обов'язково:
- використовуй веб-пошук для підтвердження фактів;
- не вигадуй дані;
- якщо факт не встановлено, пиши рівно "не встановлено";
- не змішуй активи самої компанії з активами всієї групи;
- у джерелах повертай 5-8 найрелевантніших записів з робочими URL;
- для карти повертай лише підтверджені ключові фізичні активи бізнес-групи;
- якщо точна геолокація невідома, можна вказати тільки region_id;
- beneficiaries, events, sources мають бути масивами об'єктів.

Поверни ТІЛЬКИ один JSON-об'єкт без markdown.
""".strip()


def research_company_report(edrpou: str, company_name_hint: str | None = None, additional_instructions: str | None = None, config: ReportConfig | None = None) -> dict[str, Any]:
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY is not configured")

    config = config or ReportConfig()
    client = OpenAI()
    response = client.responses.create(
        model=config.openai_model,
        reasoning={"effort": config.openai_reasoning_effort},
        tools=[{"type": config.openai_web_search_tool_type}],
        input=[
            {
                "role": "system",
                "content": (
                    "Ти аналітик українських компаній. "
                    "Формуєш лише структуровані дані для довідки по компанії за ЄДРПОУ. "
                    "Працюй обережно з належністю до групи, бенефіціарами, подіями та активами."
                ),
            },
            {
                "role": "user",
                "content": _research_prompt(edrpou, company_name_hint, additional_instructions),
            },
        ],
    )
    return _normalize_report_payload(_extract_json_object(response.output_text), edrpou)


def render_map_png(map_payload: dict[str, Any], base_dir: Path) -> bytes:
    safe_payload = {
        "group_name": _normalize_text(map_payload.get("group_name")),
        "description": _normalize_text(map_payload.get("description"), "Карта активів бізнес-групи"),
        "show_title": bool(map_payload.get("show_title", False)),
        "show_legend": bool(map_payload.get("show_legend", False)),
        "derive_active_regions": bool(map_payload.get("derive_active_regions", True)),
        "active_regions": map_payload.get("active_regions", []),
        "assets": map_payload.get("assets", []) if isinstance(map_payload.get("assets"), list) else [],
    }

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        data_path = tmp_path / "map_data.json"
        svg_path = tmp_path / "map.svg"
        png_path = tmp_path / "map.png"
        data_path.write_text(json.dumps(safe_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        render_business_map(
            data_path=data_path,
            output_svg=svg_path,
            template_path=base_dir / "ua_business_map_release" / "ua_business_map_template.svg",
            manifest_path=base_dir / "ua_business_map_release" / "ua_business_map_manifest.json",
            output_png=png_path,
        )
        return png_path.read_bytes()


def build_docx_report(report_data: dict[str, Any], map_png_bytes: bytes, template_path: Path) -> bytes:
    document = Document(str(template_path))
    company = report_data["company"]
    group = report_data["group"]

    replacements = {
        "{{COMPANY_NAME}}": company["company_name"],
        "{{EDRPOU}}": company["edrpou"],
        "{{REPORT_DATE}}": report_data["report_date"],
        "{{FULL_NAME}}": company["full_name"],
        "{{STATUS}}": company["status"],
        "{{REGISTRATION_DATE}}": company["registration_date"],
        "{{LEGAL_ADDRESS}}": company["legal_address"],
        "{{DIRECTOR}}": company["director"],
        "{{MAIN_KVED}}": company["main_kved"],
        "{{COMPANY_ACTIVITY_AND_ASSETS}}": company["activity_and_assets"],
        "{{GROUP_MEMBERSHIP_AND_ROLE}}": group["membership_and_role"],
        "{{GROUP_NAME}}": group["group_name"],
        "{{CONTROLLING_PERSONS}}": group["controlling_persons"],
        "{{GROUP_BUSINESS_LINES}}": group["business_lines"],
        "{{GROUP_GEOGRAPHY}}": group["geography"],
        "{{GROUP_ASSET_TYPES}}": group["asset_types"],
        "{{KEY_GROUP_COMPANIES}}": group["key_group_companies"],
    }
    _replace_tokens(document, replacements)
    _insert_map_image(document, map_png_bytes)

    beneficiaries_table = _find_table_after_heading(document, "5. Кінцеві бенефіціари")
    beneficiary_rows = report_data["beneficiaries"] or []
    if beneficiary_rows:
        _fill_repeating_table(
            beneficiaries_table,
            "{{BENEFICIARY_PERSON}}",
            [
                {
                    "{{BENEFICIARY_PERSON}}": _normalize_text(item.get("person")),
                    "{{BENEFICIARY_STATUS}}": _normalize_text(item.get("status"), ""),
                    "{{CONTROL_BASIS}}": _normalize_text(item.get("control_basis"), ""),
                    "{{RELIABILITY}}": _normalize_text(item.get("reliability"), ""),
                }
                for item in beneficiary_rows
            ],
        )
    else:
        _fill_repeating_table(beneficiaries_table, "{{BENEFICIARY_PERSON}}", _fallback_beneficiaries())

    events_table = _find_table_after_heading(document, "6. Істотні події та ризики")
    event_rows = report_data["events"] or []
    if event_rows:
        _fill_repeating_table(
            events_table,
            "{{EVENT_OR_RISK}}",
            [
                {
                    "{{EVENT_OR_RISK}}": _normalize_text(item.get("event_or_risk")),
                    "{{IMPACT}}": _normalize_text(item.get("impact"), ""),
                    "{{DATE_OR_PERIOD}}": _normalize_text(item.get("date_or_period"), ""),
                }
                for item in event_rows
            ],
        )
    else:
        _fill_repeating_table(events_table, "{{EVENT_OR_RISK}}", _fallback_events())

    sources_table = _find_table_after_heading(document, "7. Джерела")
    source_rows = report_data["sources"] or []
    if source_rows:
        _fill_repeating_table(
            sources_table,
            "{{SOURCE_NAME}}",
            [
                {
                    "{{SOURCE_NAME}}": _normalize_text(item.get("source_name")),
                    "{{SOURCE_DATE}}": _normalize_text(item.get("source_date"), ""),
                    "{{SOURCE_FACT}}": _normalize_text(item.get("source_fact"), ""),
                    "{{SOURCE_URL}}": "",
                }
                for item in source_rows
            ],
        )
        data_rows = sources_table.rows[1:]
        for row, source in zip(data_rows, source_rows):
            url_cell = row.cells[3]
            for paragraph in url_cell.paragraphs:
                _replace_paragraph_text(paragraph, "")
            paragraph = url_cell.paragraphs[0]
            url = _normalize_text(source.get("source_url"), "")
            if url:
                _add_hyperlink(paragraph, url, url)
            else:
                paragraph.add_run("")
    else:
        _fill_repeating_table(sources_table, "{{SOURCE_NAME}}", _fallback_sources())

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_report_filename(edrpou: str) -> str:
    return f"CompanyInfo_UA_{edrpou}_{REPORT_DATE}.docx"


def build_download_name(filename: str) -> str:
    return quote(filename)


def generate_report_bundle(
    edrpou: str,
    company_name_hint: str | None,
    additional_instructions: str | None,
    base_dir: Path,
    config: ReportConfig | None = None,
    live_mode: bool = False,
) -> dict[str, Any]:
    config = config or ReportConfig()
    if live_mode:
        if not config.allow_live_openai:
            raise RuntimeError("Live OpenAI mode is disabled. Set ALLOW_LIVE_OPENAI=true to enable paid requests.")
        report_data = research_company_report(
            edrpou=edrpou,
            company_name_hint=company_name_hint,
            additional_instructions=additional_instructions,
            config=config,
        )
    else:
        report_data = build_mock_report_data(edrpou=edrpou, company_name_hint=company_name_hint)

    map_png_bytes = render_map_png(report_data["map"], base_dir=base_dir)
    docx_bytes = build_docx_report(
        report_data=report_data,
        map_png_bytes=map_png_bytes,
        template_path=base_dir / "CompanyInfo_UA_docx_setup" / "CompanyInfo_UA_report_template.docx",
    )
    filename = build_report_filename(edrpou)
    return {
        "filename": filename,
        "download_name": build_download_name(filename),
        "report_data": report_data,
        "live_mode": live_mode,
        "docx_bytes": docx_bytes,
        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "map_png_base64": base64.b64encode(map_png_bytes).decode("ascii"),
    }
